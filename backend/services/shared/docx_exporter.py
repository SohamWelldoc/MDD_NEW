"""High-quality Markdown-to-DOCX exporter with Mermaid diagram rendering."""

from __future__ import annotations

import base64
import gzip
import os
import shutil
import subprocess
import re
import tempfile
from pathlib import Path
from typing import List, Optional, Tuple
from urllib.request import urlopen


def markdown_to_docx(
    markdown: str,
    output_path: str,
    *,
    document_title: Optional[str] = None,
    figure_captions: Optional[List[str]] = None,
) -> str:
    """Convert Markdown to DOCX, rendering Mermaid fences as images first."""
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    keep_assets = os.getenv("KEEP_DOCX_ASSETS", "false").lower() in ("1", "true", "yes")
    if keep_assets:
        work_dir = output.parent / f"{output.stem}_docx_assets"
        _write_docx_from_work_dir(
            markdown,
            output,
            work_dir,
            keep_assets=True,
            document_title=document_title,
            figure_captions=figure_captions,
        )
    else:
        with tempfile.TemporaryDirectory(prefix=f"{output.stem}_docx_") as tmp:
            _write_docx_from_work_dir(
                markdown,
                output,
                Path(tmp),
                keep_assets=False,
                document_title=document_title,
                figure_captions=figure_captions,
            )

    return str(output)


def _write_docx_from_work_dir(
    markdown: str,
    output: Path,
    work_dir: Path,
    *,
    keep_assets: bool,
    document_title: Optional[str],
    figure_captions: Optional[List[str]],
) -> None:
    diagrams_dir = work_dir / "diagrams"
    diagrams_dir.mkdir(parents=True, exist_ok=True)

    normalized = _normalize_tables_for_pandoc(markdown)
    processed_markdown, diagram_paths = _replace_mermaid_blocks(
        normalized,
        diagrams_dir,
        output.stem,
        figure_captions=figure_captions,
    )
    processed_path = work_dir / f"{output.stem}.pandoc.md"
    processed_path.write_text(processed_markdown, encoding="utf-8")

    _convert_with_pandoc(processed_path, output)
    _apply_word_branding(output, markdown, document_title=document_title)

    if keep_assets:
        manifest = work_dir / "diagrams.txt"
        manifest.write_text("\n".join(str(path) for path in diagram_paths), encoding="utf-8")


def _replace_mermaid_blocks(
    markdown: str,
    diagrams_dir: Path,
    stem: str,
    *,
    figure_captions: Optional[List[str]] = None,
) -> Tuple[str, List[Path]]:
    diagram_paths: List[Path] = []

    def replace(match: re.Match) -> str:
        index = len(diagram_paths) + 1
        body = match.group(1).strip()
        image_path = diagrams_dir / f"{stem}_diagram_{index}.png"
        _render_mermaid(body, image_path)
        diagram_paths.append(image_path)
        rel_path = image_path.as_posix()
        caption = ""
        if figure_captions and index <= len(figure_captions):
            caption = figure_captions[index - 1].strip()
        if not caption:
            caption = "Architecture Diagram"
        return f"\n\n![]({rel_path})\n\n**Figure {index}: {caption}**\n\n"

    processed = re.sub(r"```mermaid\s*\n(.*?)```", replace, markdown, flags=re.DOTALL)
    return _normalize_markdown_spacing(processed), diagram_paths


def _normalize_tables_for_pandoc(markdown: str) -> str:
    """Dedent pipe tables nested below bullet labels so Pandoc emits real DOCX tables."""
    lines = markdown.splitlines()
    out: List[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        label_match = re.match(r"^-\s+(.+?:)\s*$", line)
        if label_match and i + 2 < len(lines) and _is_indented_table_row(lines[i + 1]):
            label = label_match.group(1).strip()
            out.append(f"**{label}**")
            out.append("")
            i += 1
            while i < len(lines) and _is_indented_table_row(lines[i]):
                out.append(lines[i].strip())
                i += 1
            out.append("")
            continue
        out.append(line)
        i += 1
    return "\n".join(out)


def _is_indented_table_row(line: str) -> bool:
    stripped = line.strip()
    return line[:1].isspace() and stripped.startswith("|") and stripped.endswith("|")


def _render_mermaid(source: str, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    mmd_path = output_path.with_suffix(".mmd")
    mmd_path.write_text(source, encoding="utf-8")

    mmdc = _mermaid_cli_path()
    if mmdc:
        cmd = [str(mmdc), "-i", str(mmd_path), "-o", str(output_path), "-b", "white", "--scale", "2"]
        puppeteer_config = _puppeteer_config_path()
        if puppeteer_config:
            cmd.extend(["-p", str(puppeteer_config)])
        try:
            subprocess.run(
                cmd,
                check=True,
                capture_output=True,
                text=True,
                env=_mermaid_env(),
            )
            return
        except Exception as exc:  # noqa: BLE001
            print(f"[DOCX Export] Mermaid CLI rendering failed, trying fallback renderers: {exc}")

    _render_mermaid_with_kroki(source, output_path)


def _mermaid_cli_path() -> Path | None:
    configured = os.getenv("MERMAID_CLI_PATH")
    if configured:
        path = Path(configured)
        if not path.is_absolute():
            path = Path.cwd() / path
        if path.is_file():
            return path

    found = shutil.which("mmdc")
    if found:
        return Path(found)

    root = Path(__file__).resolve().parents[3]
    local = root / "tools" / "mermaid" / "node_modules" / ".bin" / ("mmdc.cmd" if os.name == "nt" else "mmdc")
    return local if local.is_file() else None


def _puppeteer_config_path() -> Path | None:
    configured = os.getenv("PUPPETEER_CONFIG_PATH") or os.getenv("MERMAID_PUPPETEER_CONFIG")
    if configured:
        path = Path(configured)
        if not path.is_absolute():
            path = Path.cwd() / path
        return path if path.is_file() else None

    root = Path(__file__).resolve().parents[3]
    local = root / "tools" / "puppeteer-config.json"
    return local if local.is_file() else None


def _mermaid_env() -> dict[str, str]:
    env = os.environ.copy()
    root = Path(__file__).resolve().parents[3]
    node_dir = root / "tools" / "node"
    if node_dir.is_dir():
        env["PATH"] = str(node_dir) + os.pathsep + env.get("PATH", "")
    return env


def _render_mermaid_with_kroki(source: str, output_path: Path) -> None:
    encoded = base64.urlsafe_b64encode(gzip.compress(source.encode("utf-8"), compresslevel=9)).decode("ascii")
    url = f"https://kroki.io/mermaid/png/{encoded}"
    try:
        with urlopen(url, timeout=30) as response:
            output_path.write_bytes(response.read())
    except Exception as exc:  # noqa: BLE001
        print(f"[DOCX Export] Kroki Mermaid rendering failed, using local fallback: {exc}")
        _render_mermaid_locally(source, output_path)


def _render_mermaid_locally(source: str, output_path: Path) -> None:
    stripped = source.strip()
    if stripped.startswith("sequenceDiagram"):
        _render_sequence_diagram(stripped, output_path)
        return
    if stripped.startswith("flowchart") or stripped.startswith("graph"):
        _render_flowchart(stripped, output_path)
        return
    _render_mermaid_source_image(stripped, output_path)


def _render_flowchart(source: str, output_path: Path) -> None:
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch

    nodes: dict[str, str] = {}
    edges: list[tuple[str, str, str]] = []
    for raw_line in source.splitlines()[1:]:
        line = raw_line.strip().rstrip(";")
        if not line or line.startswith("%%"):
            continue
        edge_match = re.search(r"([A-Za-z0-9_]+).*?--(?:>|>|-)?(?:\|\"?([^|]+?)\"?\|)?\s*([A-Za-z0-9_]+)", line)
        if edge_match:
            left, label, right = edge_match.groups()
            nodes.setdefault(left, _extract_node_label(line, left))
            nodes.setdefault(right, _extract_node_label(line, right))
            edges.append((left, right, label or ""))
            continue
        node_match = re.match(r"([A-Za-z0-9_]+)(?:\[(.+?)\]|\((.+?)\)|\{\"?(.+?)\"?\})", line)
        if node_match:
            node_id = node_match.group(1)
            nodes[node_id] = _clean_label(next(group for group in node_match.groups()[1:] if group))

    if not nodes:
        _render_mermaid_source_image(source, output_path)
        return

    levels: dict[str, int] = {node: 0 for node in nodes}
    for _ in range(len(nodes)):
        changed = False
        for src, dst, _label in edges:
            new_level = levels.get(src, 0) + 1
            if new_level > levels.get(dst, 0):
                levels[dst] = new_level
                changed = True
        if not changed:
            break

    by_level: dict[int, list[str]] = {}
    for node, level in levels.items():
        by_level.setdefault(level, []).append(node)

    positions: dict[str, tuple[float, float]] = {}
    max_width = max(len(items) for items in by_level.values())
    for level, items in sorted(by_level.items()):
        for idx, node in enumerate(items):
            x = (idx - (len(items) - 1) / 2) * 5.0
            y = -level * 3.0
            positions[node] = (x, y)

    fig_width = max(9, max_width * 4.4)
    fig_height = max(5.6, (max(by_level) + 1) * 2.8)
    fig, ax = plt.subplots(figsize=(fig_width, fig_height), dpi=180)
    ax.axis("off")

    for src, dst, label in edges:
        if src not in positions or dst not in positions:
            continue
        x1, y1 = positions[src]
        x2, y2 = positions[dst]
        start, end = _edge_points((x1, y1), (x2, y2), half_width=1.5, half_height=0.58)
        annotation = ax.annotate(
            "",
            xy=end,
            xytext=start,
            arrowprops=dict(arrowstyle="->", linewidth=1.5, color="#344054", shrinkA=0, shrinkB=0),
        )
        annotation.set_clip_on(False)
        if label:
            ax.text(
                (start[0] + end[0]) / 2,
                (start[1] + end[1]) / 2 + 0.32,
                _wrap_label(_clean_label(label), 28),
                ha="center",
                va="bottom",
                fontsize=8,
                bbox=dict(facecolor="white", edgecolor="none", pad=0.4),
                clip_on=False,
            )

    for node, (x, y) in positions.items():
        label = _wrap_label(nodes[node], width=24)
        rect = FancyBboxPatch(
            (x - 1.35, y - 0.45),
            2.7,
            0.9,
            boxstyle="round,pad=0.12,rounding_size=0.08",
            linewidth=1.3,
            edgecolor="#1d4ed8",
            facecolor="#eff6ff",
        )
        ax.add_patch(rect)
        ax.text(x, y, label, ha="center", va="center", fontsize=8.5, color="#111827", clip_on=False)

    xs = [p[0] for p in positions.values()]
    ys = [p[1] for p in positions.values()]
    ax.set_xlim(min(xs) - 3.1, max(xs) + 3.1)
    ax.set_ylim(min(ys) - 1.8, max(ys) + 1.8)
    fig.tight_layout(pad=1.2)
    fig.savefig(output_path, bbox_inches="tight", pad_inches=0.35, facecolor="white")
    plt.close(fig)


def _edge_points(
    start: tuple[float, float],
    end: tuple[float, float],
    *,
    half_width: float,
    half_height: float,
) -> tuple[tuple[float, float], tuple[float, float]]:
    """Return line endpoints that stop at rectangular node edges."""
    x1, y1 = start
    x2, y2 = end
    dx = x2 - x1
    dy = y2 - y1
    if dx == 0 and dy == 0:
        return start, end

    if abs(dx) / half_width > abs(dy) / half_height:
        sx = half_width if dx > 0 else -half_width
        sy = dy * (abs(sx) / abs(dx)) if dx else 0
    else:
        sy = half_height if dy > 0 else -half_height
        sx = dx * (abs(sy) / abs(dy)) if dy else 0

    return (x1 + sx, y1 + sy), (x2 - sx, y2 - sy)


def _render_sequence_diagram(source: str, output_path: Path) -> None:
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch

    participants: list[str] = []
    aliases: dict[str, str] = {}
    messages: list[tuple[str, str, str]] = []
    for raw_line in source.splitlines()[1:]:
        line = raw_line.strip()
        if not line or line.startswith("%%") or line.startswith("Note "):
            continue
        participant = re.match(r"participant\s+([A-Za-z0-9_]+)(?:\s+as\s+\"?(.+?)\"?)?$", line)
        if participant:
            pid, label = participant.groups()
            if pid not in participants:
                participants.append(pid)
            aliases[pid] = _clean_label(label or pid)
            continue
        msg = re.match(r"([A-Za-z0-9_]+)\s*-+>>?\s*([A-Za-z0-9_]+)\s*:\s*(.+)$", line)
        if msg:
            src, dst, label = msg.groups()
            for pid in (src, dst):
                if pid not in participants:
                    participants.append(pid)
                    aliases[pid] = pid
            messages.append((src, dst, _clean_label(label)))

    if not participants:
        _render_mermaid_source_image(source, output_path)
        return

    width = max(9, len(participants) * 2.55)
    height = max(5.8, len(messages) * 0.9 + 2.8)
    fig, ax = plt.subplots(figsize=(width, height), dpi=180)
    ax.axis("off")

    x_positions = {pid: idx * 2.7 for idx, pid in enumerate(participants)}
    top_y = 0
    bottom_y = -(len(messages) + 1) * 0.9

    for pid in participants:
        x = x_positions[pid]
        rect = FancyBboxPatch(
            (x - 0.8, top_y - 0.35),
            1.6,
            0.7,
            boxstyle="round,pad=0.08",
            linewidth=1.2,
            edgecolor="#7c3aed",
            facecolor="#f5f3ff",
        )
        ax.add_patch(rect)
        ax.text(x, top_y, _wrap_label(aliases.get(pid, pid), 16), ha="center", va="center", fontsize=8, clip_on=False)
        ax.plot([x, x], [top_y - 0.45, bottom_y], linestyle="--", color="#98a2b3", linewidth=1)

    for idx, (src, dst, label) in enumerate(messages, start=1):
        y = -idx * 0.9
        x1 = x_positions[src]
        x2 = x_positions[dst]
        direction = 1 if x2 >= x1 else -1
        start_x = x1 + direction * 0.18
        end_x = x2 - direction * 0.18
        annotation = ax.annotate(
            "",
            xy=(end_x, y),
            xytext=(start_x, y),
            arrowprops=dict(arrowstyle="->", linewidth=1.3, color="#344054", shrinkA=0, shrinkB=0),
        )
        annotation.set_clip_on(False)
        ax.text(
            (start_x + end_x) / 2,
            y + 0.26,
            _wrap_label(label, 34),
            ha="center",
            va="bottom",
            fontsize=7.5,
            bbox=dict(facecolor="white", edgecolor="none", pad=0.35),
            clip_on=False,
        )

    ax.set_xlim(-1.4, max(x_positions.values()) + 1.4)
    ax.set_ylim(bottom_y - 0.8, top_y + 1.0)
    fig.tight_layout(pad=1.2)
    fig.savefig(output_path, bbox_inches="tight", pad_inches=0.35, facecolor="white")
    plt.close(fig)


def _render_mermaid_source_image(source: str, output_path: Path) -> None:
    import matplotlib.pyplot as plt

    lines = source.splitlines()
    height = max(3, min(12, len(lines) * 0.28 + 1))
    fig, ax = plt.subplots(figsize=(10, height), dpi=180)
    ax.axis("off")
    ax.text(
        0.01,
        0.99,
        source,
        ha="left",
        va="top",
        family="monospace",
        fontsize=8,
        bbox=dict(boxstyle="round,pad=0.5", facecolor="#f8fafc", edgecolor="#cbd5e1"),
    )
    fig.tight_layout(pad=1.1)
    fig.savefig(output_path, bbox_inches="tight", pad_inches=0.25, facecolor="white")
    plt.close(fig)


def _extract_node_label(line: str, node_id: str) -> str:
    match = re.search(rf"{re.escape(node_id)}(?:\[(.+?)\]|\((.+?)\)|\{{\"?(.+?)\"?\}})", line)
    if not match:
        return node_id
    return _clean_label(next(group for group in match.groups() if group))


def _clean_label(value: str) -> str:
    value = (value or "").strip().strip('"')
    value = re.sub(r"<br\s*/?>", "\n", value, flags=re.IGNORECASE)
    value = value.replace("&nbsp;", " ")
    return value


def _wrap_label(value: str, width: int) -> str:
    words = value.split()
    if not words:
        return value
    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        if len(current) + len(word) + 1 > width:
            lines.append(current)
            current = word
        else:
            current += " " + word
    lines.append(current)
    return "\n".join(lines)


def _convert_with_pandoc(markdown_path: Path, output_path: Path) -> None:
    try:
        import pypandoc
    except ImportError as exc:  # pragma: no cover - depends on environment
        raise RuntimeError("pypandoc_binary is required for DOCX export. Install requirements.txt.") from exc

    extra_args = [
        "--from=markdown+pipe_tables+fenced_code_blocks+backtick_code_blocks",
        "--standalone",
        "--resource-path",
        str(markdown_path.parent),
    ]
    reference_doc = _reference_docx()
    if reference_doc:
        extra_args.extend(["--reference-doc", str(reference_doc)])

    pypandoc.convert_file(
        str(markdown_path),
        "docx",
        outputfile=str(output_path),
        extra_args=extra_args,
    )


def _reference_docx() -> Path | None:
    import os

    value = os.getenv("DOCX_REFERENCE_PATH") or os.getenv("MDD_TEMPLATE_PATH")
    if not value:
        return None
    path = Path(value)
    if not path.is_absolute():
        path = Path.cwd() / path
    return path if path.is_file() else None


_WELLDOC_FOOTER_TEXT = (
    "© 2009-24 Welldoc, Inc. Intellectual Property. All rights reserved. Proprietary and Confidential. "
    "Not permitted to be duplicated or reproduced without the express written consent of Welldoc, Inc. "
    "Welldoc and the logos associated therewith and all other Welldoc marks contained herein are trademarks "
    "of Welldoc. All other marks contained herein are the property of their respective owners"
)


def _apply_word_branding(docx_path: Path, markdown: str, *, document_title: Optional[str] = None) -> None:
    """Apply Welldoc-ready page furniture and readable Word formatting."""
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION_START
        from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:  # pragma: no cover - depends on environment
        raise RuntimeError("python-docx is required for branded DOCX export. Install requirements.txt.") from exc

    doc = Document(str(docx_path))
    header_logo = _header_logo_path()
    resolved_document_title = document_title or _document_title(markdown)
    document_id = os.getenv("HLD_DOCUMENT_ID", "")

    for section in doc.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.header_distance = Inches(0.12)
        section.footer_distance = Inches(0.16)
        section.different_first_page_header_footer = False

        _clear_container(section.header)
        _add_dynamic_header(
            section.header,
            header_logo=header_logo,
            title=resolved_document_title,
            document_id=document_id,
            Inches=Inches,
            Pt=Pt,
            RGBColor=RGBColor,
            WD_ALIGN_VERTICAL=WD_ALIGN_VERTICAL,
            OxmlElement=OxmlElement,
            qn=qn,
        )

        _clear_container(section.footer)
        paragraph = section.footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(_WELLDOC_FOOTER_TEXT)
        _set_run_font(run, "Times New Roman", Pt(7.5), RGBColor(0, 0, 0))

    _style_document_styles(doc, Pt, RGBColor)
    _style_body_paragraphs(doc, Pt, RGBColor, WD_ALIGN_PARAGRAPH)
    _style_tables(doc, Pt, RGBColor, WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, OxmlElement, qn)
    _style_images(doc, Inches, WD_ALIGN_PARAGRAPH)

    doc.save(str(docx_path))


def _style_document_styles(doc, Pt, RGBColor) -> None:
    for style_name in ("Normal", "Body Text"):
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style.font.size = Pt(11)
            style.font.color.rgb = RGBColor(0, 0, 0)

    for style in doc.styles:
        if style.name.startswith("Heading"):
            style.font.name = "Times New Roman"
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 0, 0)


def _add_dynamic_header(
    header,
    *,
    header_logo: Path | None,
    title: str,
    document_id: str,
    Inches,
    Pt,
    RGBColor,
    WD_ALIGN_VERTICAL,
    OxmlElement,
    qn,
) -> None:
    table = header.add_table(rows=3, cols=2, width=Inches(6.9))
    table.autofit = False
    _set_table_borders(table, OxmlElement, qn)

    logo_cell = table.cell(0, 0)
    for row in table.rows[1:]:
        logo_cell = logo_cell.merge(row.cells[0])
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    logo_paragraph = logo_cell.paragraphs[0]
    logo_paragraph.paragraph_format.space_after = Pt(0)
    if header_logo:
        logo_paragraph.add_run().add_picture(str(header_logo), width=Inches(2.45))

    right_labels = [
        ("TITLE: ", title),
        ("Document ID: ", document_id),
        ("PAGE: Page ", None),
    ]
    for row_index, (label, value) in enumerate(right_labels):
        cell = table.cell(row_index, 1)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(label)
        _set_run_font(run, "Times New Roman", Pt(9), RGBColor(0, 0, 0))
        if value is None:
            _add_field(paragraph, "PAGE", "Times New Roman", Pt(9), RGBColor(0, 0, 0))
            run = paragraph.add_run(" of ")
            _set_run_font(run, "Times New Roman", Pt(9), RGBColor(0, 0, 0))
            _add_field(paragraph, "NUMPAGES", "Times New Roman", Pt(9), RGBColor(0, 0, 0))
        elif value:
            value_run = paragraph.add_run(value)
            _set_run_font(value_run, "Times New Roman", Pt(9), RGBColor(0, 0, 0))

    for row in table.rows:
        row.cells[0].width = Inches(2.75)
        row.cells[1].width = Inches(4.15)

    trailing = header.add_paragraph()
    trailing.paragraph_format.space_after = Pt(0)


def _add_field(paragraph, field_code: str, font_name: str, font_size, font_color) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    _set_run_font(run, font_name, font_size, font_color)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(text)
    run._r.append(fld_end)


def _style_body_paragraphs(doc, Pt, RGBColor, WD_ALIGN_PARAGRAPH) -> None:
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        fmt = paragraph.paragraph_format
        fmt.line_spacing = 1.0

        if style_name.startswith("Heading"):
            fmt.space_before = Pt(12 if style_name == "Heading 1" else 6)
            fmt.space_after = Pt(24 if style_name == "Heading 1" else 12)
        elif _is_figure_caption(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt.space_before = Pt(2)
            fmt.space_after = Pt(12)
        else:
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(6)

        for run in paragraph.runs:
            is_heading = style_name.startswith("Heading")
            is_caption = _is_figure_caption(paragraph)
            _set_run_font(run, "Times New Roman", Pt(11), RGBColor(0, 0, 0), bold=is_heading or is_caption)


def _style_tables(doc, Pt, RGBColor, WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, OxmlElement, qn) -> None:
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        _set_table_borders(table, OxmlElement, qn)

        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                _set_cell_shading(cell, "F2F2F2" if row_index == 0 else "FFFFFF", OxmlElement, qn)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(2)
                    for run in paragraph.runs:
                        _set_run_font(
                            run,
                            "Times New Roman",
                            Pt(10),
                            RGBColor(0, 0, 0),
                            bold=row_index == 0,
                        )


def _style_images(doc, Inches, WD_ALIGN_PARAGRAPH) -> None:
    for shape in doc.inline_shapes:
        if shape.width and shape.width > Inches(1.0):
            max_width = Inches(6.15)
            max_height = Inches(4.25)
            if shape.width != max_width:
                ratio = shape.height / shape.width if shape.width else 1
                shape.width = max_width
                shape.height = int(max_width * ratio)
            if shape.height > max_height:
                ratio = shape.width / shape.height if shape.height else 1
                shape.height = max_height
                shape.width = int(max_height * ratio)

    for paragraph in doc.paragraphs:
        if _paragraph_has_image(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Inches(0)
            paragraph.paragraph_format.space_after = Inches(0.08)


def _set_run_font(run, name, size, color, *, bold: bool | None = None) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run.font.name = name
    run.font.size = size
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)
    r_fonts.set(qn("w:cs"), name)


def _set_table_borders(table, OxmlElement, qn) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def _set_cell_shading(cell, fill: str, OxmlElement, qn) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _clear_container(container) -> None:
    for child in list(container._element):
        container._element.remove(child)


def _is_figure_caption(paragraph) -> bool:
    return paragraph.text.strip().lower().startswith("figure ")


def _paragraph_has_image(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def _header_logo_path() -> Path | None:
    value = os.getenv("HLD_LOGO_PATH") or os.getenv("DOCX_LOGO_PATH")
    if value:
        path = Path(value)
        if not path.is_absolute():
            path = Path.cwd() / path
        return path if path.is_file() else None

    default = Path(__file__).resolve().parents[2] / "assets" / "welldoc_logo.png"
    return default if default.is_file() else None


def _document_title(markdown: str) -> str:
    explicit = os.getenv("HLD_DOCUMENT_TITLE") or os.getenv("DOCX_DOCUMENT_TITLE")
    if explicit:
        return explicit

    topic = _topic_from_scope(markdown) or _topic_from_heading(markdown) or "System"
    topic = re.sub(r"\s+feature$", "", topic.strip(), flags=re.IGNORECASE)
    topic = re.sub(r"\s+high[- ]level\s+design$", "", topic.strip(), flags=re.IGNORECASE)
    return f"{topic} High Level Design"


def _topic_from_scope(markdown: str) -> str | None:
    match = re.search(
        r"The in-scope features/modules are:\s*\n(?:\*|-)\s+(.+?)(?:\s+feature)?\s*(?:\n|$)",
        markdown,
        flags=re.IGNORECASE,
    )
    if not match:
        return None
    return _clean_label(match.group(1)).strip()


def _topic_from_heading(markdown: str) -> str | None:
    for line in markdown.splitlines():
        match = re.match(r"^#\s+(.+?)\s*$", line)
        if match:
            heading = match.group(1).strip()
            heading = re.split(r"\s+[—-]\s+High[- ]Level Design", heading, maxsplit=1, flags=re.IGNORECASE)[0]
            return _clean_label(heading)
    return None


def _normalize_markdown_spacing(markdown: str) -> str:
    markdown = re.sub(r"\n{4,}", "\n\n\n", markdown)
    return markdown.strip() + "\n"
